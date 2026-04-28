import io
import base64
import asyncio
import logging

import pdfplumber
import fitz  # PyMuPDF
import docx

from config import env_int

logger = logging.getLogger(__name__)

MAX_TEXT_LENGTH = env_int("CV_TEXT_MAX_CHARS", 12000)
MAX_VISION_PAGES = env_int("CV_VISION_MAX_PAGES", 3)

# En dessous de ce seuil on considère le texte trop pauvre pour une analyse fiable
_OCR_FALLBACK_MIN_CHARS = 80
# Timeout total d'extraction par fichier (secondes)
_EXTRACT_TIMEOUT_S = 30

SUPPORTED_EXTENSIONS = {"pdf", "docx"}

# Magic bytes that identify genuine PDF and DOCX (ZIP-based) containers.
_PDF_MAGIC = b"%PDF"
_DOCX_MAGIC = b"PK\x03\x04"


def validate_file_magic(content: bytes, ext: str) -> bool:
    """Returns True only when the file's leading bytes match the declared extension.

    Prevents attackers from uploading executables or malicious payloads renamed
    to .pdf / .docx. Must be called before any parsing logic.
    """
    if ext == "pdf":
        return content[:4] == _PDF_MAGIC
    if ext == "docx":
        return content[:4] == _DOCX_MAGIC
    return False


def _extract_pdf(content: bytes) -> str:
    """Extraction texte via pdfplumber (PDFs avec couche texte)."""
    try:
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            pages = []
            for page in pdf.pages[:8]:
                text = page.extract_text()
                if text:
                    pages.append(text)
            return "\n".join(pages)[:MAX_TEXT_LENGTH]
    except Exception as e:
        logger.warning("pdfplumber failed: %s", e)
        return ""


def _pdf_pages_as_jpeg_b64(content: bytes, max_pages: int = 3) -> list:
    """Convertit les premières pages d'un PDF en JPEG base64 via PyMuPDF (rapide, sans tesseract)."""
    try:
        doc = fitz.open(stream=content, filetype="pdf")
        results = []
        for i in range(min(max_pages, len(doc))):
            page = doc[i]
            # ~150 DPI : bon pour Claude Vision, taille raisonnable
            mat = fitz.Matrix(150 / 72, 150 / 72)
            pix = page.get_pixmap(matrix=mat, colorspace=fitz.csGRAY, alpha=False)
            jpeg_bytes = pix.tobytes("jpeg", jpg_quality=85)
            results.append(base64.b64encode(jpeg_bytes).decode())
        return results
    except Exception as e:
        logger.error("PyMuPDF image extraction failed: %s", e)
        return []


def _extract_docx(content: bytes) -> str:
    """Extraction Word (.docx) texte + tableaux."""
    try:
        doc = docx.Document(io.BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)[:MAX_TEXT_LENGTH]
    except Exception as e:
        logger.error("DOCX extraction failed: %s", e)
        raise ValueError(f"Erreur extraction Word: {e}") from e


async def extract_text(filename: str, content: bytes) -> tuple:
    """Extraction depuis un PDF ou Word.

    Retourne (text, images_b64) où images_b64 est une liste non vide uniquement
    si le texte est insuffisant et que des images ont pu être extraites (PDFs scannés).
    """
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""

    if ext == "pdf":
        try:
            text = await asyncio.wait_for(
                asyncio.to_thread(_extract_pdf, content),
                timeout=_EXTRACT_TIMEOUT_S,
            )
        except asyncio.TimeoutError:
            logger.error("pdfplumber timeout pour %s", filename)
            text = ""

        if len(text.strip()) >= _OCR_FALLBACK_MIN_CHARS:
            return text, []

        # Texte insuffisant : extraire les images pour Claude Vision
        logger.info("Texte insuffisant (%d chars), extraction images pour Claude Vision", len(text.strip()))
        try:
            images = await asyncio.wait_for(
                asyncio.to_thread(_pdf_pages_as_jpeg_b64, content, MAX_VISION_PAGES),
                timeout=20,
            )
        except asyncio.TimeoutError:
            logger.error("PyMuPDF timeout pour %s", filename)
            images = []

        return text, images

    if ext == "docx":
        text = await asyncio.to_thread(_extract_docx, content)
        return text, []

    raise ValueError(
        f"Format non supporté pour « {filename} » : seuls les fichiers PDF (.pdf) et Word (.docx) sont acceptés."
    )
